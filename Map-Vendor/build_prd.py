from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "地图商家预约系统_后端与数据库PRD_v1.7.docx"
FONT = "Microsoft YaHei"
NAVY = "163A5F"; BLUE = "276FBF"; LIGHT = "EAF2F8"; GRAY = "667085"; GRID = "CBD5E1"; RED = "A61B1B"; GOLD = "8A5A00"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

def font(run, size=10.5, bold=False, color="222222", italic=False):
    run.font.name = FONT; run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    for k in ("ascii", "hAnsi", "eastAsia"): rf.set(qn("w:"+k), FONT)

styles = doc.styles
normal = styles["Normal"]; normal.font.name = FONT; normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.12
for name,size,color,before,after in [("Title",25,NAVY,0,6),("Subtitle",12,GRAY,0,14),("Heading 1",16,BLUE,16,7),("Heading 2",13,BLUE,12,5),("Heading 3",11.5,NAVY,8,4)]:
    s=styles[name]; s.font.name=FONT; s.font.size=Pt(size); s.font.bold=name!="Subtitle"; s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

for sty in ("List Bullet","List Number"):
    s=styles[sty]; s.font.name=FONT; s.font.size=Pt(10.5); s.paragraph_format.left_indent=Inches(.5); s.paragraph_format.first_line_indent=Inches(-.25); s.paragraph_format.space_after=Pt(4)
    s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tc=cell._tc.get_or_add_tcPr(); m=tc.first_child_found_in("w:tcMar")
    if m is None: m=OxmlElement("w:tcMar"); tc.append(m)
    for side,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        x=m.find(qn("w:"+side))
        if x is None: x=OxmlElement("w:"+side); m.append(x)
        x.set(qn("w:w"),str(val)); x.set(qn("w:type"),"dxa")

def set_table_geometry(t, widths):
    t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    tblPr=t._tbl.tblPr
    w=tblPr.first_child_found_in("w:tblW"); w.set(qn("w:w"),str(sum(widths))); w.set(qn("w:type"),"dxa")
    ind=tblPr.first_child_found_in("w:tblInd")
    if ind is None: ind=OxmlElement("w:tblInd"); tblPr.append(ind)
    ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa")
    grid=t._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for x in widths:
        gc=OxmlElement("w:gridCol"); gc.set(qn("w:w"),str(x)); grid.append(gc)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            c.width=Inches(widths[i]/1440); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(c)
            tcw=c._tc.get_or_add_tcPr().first_child_found_in("w:tcW"); tcw.set(qn("w:w"),str(widths[i])); tcw.set(qn("w:type"),"dxa")

def table(headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,LIGHT); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(h),9.1,True,NAVY)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(str(v)),8.7)
    set_table_geometry(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1); return t

def p(text="", bold=False, color="222222", size=10.5, style=None):
    q=doc.add_paragraph(style=style); font(q.add_run(text),size,bold,color); return q
def bullets(items):
    for x in items: p(x,style="List Bullet")
def nums(items):
    numbering=doc.part.numbering_part.element
    style_num_id=styles["List Number"].element.pPr.numPr.numId.val
    base_num=next(n for n in numbering.findall(qn("w:num")) if int(n.get(qn("w:numId"))) == int(style_num_id))
    abstract_id=base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id=max(int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num")))+1
    num=OxmlElement("w:num"); num.set(qn("w:numId"),str(new_num_id))
    abstract=OxmlElement("w:abstractNumId"); abstract.set(qn("w:val"),abstract_id); num.append(abstract)
    override=OxmlElement("w:lvlOverride"); override.set(qn("w:ilvl"),"0")
    start=OxmlElement("w:startOverride"); start.set(qn("w:val"),"1"); override.append(start); num.append(override)
    numbering.append(num)
    for x in items:
        q=p(x,style="List Number"); num_pr=q._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val=0; num_pr.get_or_add_numId().val=new_num_id
def h(text,l=1): doc.add_heading(text,level=l)
def callout(label,text,color=BLUE):
    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"; c=t.cell(0,0); shade(c,"F5F8FC"); margins(c,120,160,120,160)
    q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(0); font(q.add_run(label+"："),10,True,color); font(q.add_run(text),10)
    set_table_geometry(t,[9360]); doc.add_paragraph().paragraph_format.space_after=Pt(1)

# Header/footer
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(hp.add_run("地图商家预约系统｜后端与数据库 PRD"),8.5,False,GRAY)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("内部开发基线 · v1.7 修订 · 2026-08-21"),8.5,False,GRAY)

# Cover
p("产品需求文档（PRD）",True,BLUE,11)
p("地图商家预约系统",True,NAVY,25,style="Title")
p("后端开发与数据库实施版",False,GRAY,14,style="Subtitle")
table(["文档属性","内容"],[
    ["适用范围","Java Spring Boot 后端、MySQL 9.0.1、文件存储、OpenAPI、测试、部署与联调交付"],
    ["主要读者","后端开发、数据库设计、接口联调、测试与运维；前端负责人作为契约使用方"],
    ["需求来源","《地图商家预约小程序开发文档》；本文不把源文档中的建议当作用户新增指令"],
    ["版本/状态","v1.7 修订 / 小程序静默身份交换，管理员端仍为后端零登录"],
    ["编制日期","2026-08-21"],
],[1800,7560])
callout("结论","系统保持无账号密码、无 Session、无 JWT。微信小程序由 uni-app 静默取得临时 code，Java 后端仅调用微信 code2Session 并返回 openid；管理员后台仍不提供后端登录接口。")
doc.add_page_break()

h("1. 文档目的与需求边界")
h("1.1 目标",2)
p("将源开发文档转化为可开发、可联调、可验收的后端与数据库基线，明确业务范围、接口职责、数据模型、非功能要求、工程结构、开发顺序和风险。")
h("1.2 MVP 范围",2)
bullets(["高校区域、三类服务商及车辆/菜品/房型的后端 CRUD、启停、逻辑删除与图片接口。","小程序地图数据、分类计数和服务商详情查询接口。","uni-app 静默调用微信登录取得临时 code；Java 后端仅交换并返回 openid，不建立 Session/JWT。","三类预约创建、本人订单查询及待确认取消接口。","管理员订单查询、确认、取消、统计和导出接口。","统一响应、异常、参数校验、日志、数据库迁移与部署配置。","向前端负责人交付 OpenAPI、字段映射、环境地址与联调说明。"])
h("1.3 已有前端现状与协作边界",2)
table(["项目","已完成内容","当前状态/后端动作"],[
    ["管理员后台","Vue 3 + Vite + Element Plus + Pinia；已有登录、仪表盘、服务商、高校和订单页面","当前使用 mock.js 内存假数据；后端按页面用例提供真实 API，但不修改页面"],
    ["小程序","uni-app 已有首页地图和服务商详情等部分页面","由其他开发人员继续完成；后端仅交付 App API 契约"],
    ["身份入口","管理员后台已有前端演示门禁；小程序启动时静默 uni.login","后端只提供 code2Session 到 openid 的交换；不提供账号、密码、Session 或 JWT"],
],[1500,3860,4000])
bullets(["现有管理员页面仍使用 travel/hotel/food、小写状态和中文订单状态，后端契约使用 TRAVEL/HOTEL/FOOD 与 PENDING/CONFIRMED/CANCELLED；必须提供明确映射。","现有 mock 包含 price、unit、amount、酒店 area/breakfast 等字段，与当前后端口径不完全一致；未获用户确认前不得为迎合 mock 擅自新增数据库字段。","每个后端窗口以 OpenAPI、字段映射、示例请求响应和自动化测试作为前端交接物；前端联调由其他开发人员执行。"])
h("1.4 明确不在 MVP",2)
bullets(["在线支付、退款、库存锁定与价格结算。","司机端、评价体系、微信服务通知。","后台地图框选高校边界、地址自动解析经纬度。","自动确认、自动取消、资源实时可用性校验。"])
h("1.5 成功标准",2)
table(["维度","验收口径"],[
    ["功能","MVP 接口通过约定用例；待确认、已确认、已取消三种订单状态流转符合规则。"],
    ["安全","小程序静默交换不建立服务端登录态；订单接口仍信任客户端 openid。/admin/** 无后端认证，须通过 Nginx 网络边界、HTTPS、限流和审计降低风险。"],
    ["数据","核心表有约束、索引、审计字段；删除策略不会产生孤儿记录。"],
    ["性能","常规列表 P95 ≤ 500ms，地图聚合接口 P95 ≤ 800ms（单机目标，需以压测数据确认）。"],
    ["可运维","配置外置、可迁移建库、日志可追踪、可备份恢复、健康检查可用。"],
],[1500,7860])

h("2. 角色、权限与关键流程")
table(["角色","身份方式","允许操作","禁止操作"],[
    ["微信用户","uni-app 静默取 code，后端交换并返回 openid；无 Session/JWT","浏览启用内容；创建、查看、取消该 openid 的待确认预约","取消已确认订单；处理订单"],
    ["平台管理员","登录由前端自行完成；后端无账号与认证","高校、全部服务商及附属资源、文件、全部订单、统计和导出","通过业务接口删除订单；绕过审计直接改数据"],
],[1300,2100,3400,2560])
h("2.1 内容发布流程",2)
nums(["平台管理员创建并维护服务商及车辆、房型或菜品。","平台管理员控制服务商启用状态；查询端仅返回启用且未删除记录。","小程序重新请求地图接口后展示高校多边形、服务商点位与分类数量。"])
h("2.2 预约状态机",2)
table(["当前状态","动作/操作者","目标状态","约束"],[
    ["待确认","确认 / 平台管理员","已确认","仅允许一次；并发条件更新"],
    ["待确认","取消 / 平台管理员","已取消","取消原因必填，记录 ADMIN 来源"],
    ["待确认","取消 / 订单所属用户","已取消","必须验证订单归属，cancel_source=USER"],
    ["已确认/已取消","任意状态动作","不变并报业务错误","终态不可回退或二次变更"],
],[1600,2300,1700,3760])
callout("已确认规则","用户下单后为待确认，可主动取消；平台管理员确认后变为已确认，用户不可取消。管理员不接受预约时统一转为已取消，并记录 ADMIN 来源和原因。",GOLD)

h("3. 功能需求拆解")
h("3.1 登录边界",2)
table(["编号","需求","后端规则/验收"],[
    ["BOUND-01","管理员后台登录","完全由前端完成；后端不提供管理员账号、登录、改密、密码校验、Session、JWT 或令牌。"],
    ["BOUND-02","小程序静默身份","uni-app 调用 uni.login；后端调用 code2Session，仅返回 openid，不返回 session_key、不签发令牌。"],
    ["BOUND-03","管理员接口边界","/admin/** 当前不做后端登录认证；部署必须配置 Nginx 网络访问限制、HTTPS、限流和审计。"],
    ["BOUND-04","角色边界","系统不存在商家账号、商家登录、商家后台或 /merchant/**。"],
],[1000,2300,6060])
h("3.2 高校区域管理",2)
bullets(["支持列表、详情、新增、编辑、启停、预览与删除。","polygon_points 必须是合法 JSON 数组；每点含 latitude/longitude；至少 3 个不同点；经纬度范围合法。","建议数据库使用 JSON 类型；若后续需要空间查询，再升级 MySQL POLYGON/SRID 4326。","图片不使用逗号拼接，统一进入资源表或 JSON 数组；返回稳定排序。","被删除数据采用逻辑删除；启停与删除均影响小程序可见性。"])
h("3.3 服务商与附属数据",2)
table(["模块","公共字段","专属字段/子资源","关键校验"],[
    ["服务商","名称、地址、经纬度、类型、简介、状态、图片","—","经纬度范围；类型创建后原则上不可改；名称/地址长度"],
    ["出行","同上","车辆：车型、核载、图片、备注","核载 > 0；只能挂在出行服务商"],
    ["餐饮","同上","联系人、联系电话、推荐菜说明；菜品列表","菜品只能挂在餐饮服务商；MVP 不增加人均价或菜系字段"],
    ["住宿","同上","房型：名称、床型、简介、图片","房型只能挂在住宿服务商"],
],[1100,2500,3100,2660])
h("3.4 小程序查询",2)
bullets(["地图聚合接口一次返回启用高校的简要区域数据、启用服务商轻量点位、三类计数；避免 N+1 查询与传输完整详情。","支持 type 分类筛选；高校区域不因 type 改变。","服务商详情按类型返回稳定结构的 common + detail；禁用/删除服务商对小程序返回 404 或业务不可用。","列表和详情不得暴露内部审计字段、管理员信息、文件物理路径。"])
h("3.5 预约订单",2)
table(["类型","结构化公共字段","类型专属字段"],[
    ["三类公用","联系人姓名、联系电话、人数、服务日期、服务商 ID/名称快照、服务类别","user_id、订单号、状态、备注、创建/更新时间"],
    ["出行","同上","车辆规格（car_id + 名称快照）、车辆数量、服务方式（包天/往返）"],
    ["住宿","同上","房间规格（room_id + 名称快照）、房间数量"],
    ["餐饮","同上","用餐时间：早餐/中餐/晚餐，不设置具体时段"],
],[1300,3900,4160])
bullets(["创建订单时接收并保存前端传入的 openid，同时验证服务商启用、附属资源存在且归属正确、日期合法。","必须保存服务商名称、类型、所选车辆/房型等展示快照，避免管理员编辑后历史订单含义变化。","订单号由服务端生成并唯一；创建接口支持幂等键，防止重复点击产生重复订单。","个人订单按请求传入的 openid 查询；手机号不能作为订单归属凭证。","管理员订单接口可查询全部订单，并按服务日期、状态、类型和服务商筛选。","后台筛选的服务日期必须来自结构化列，不能仅解析 extra_json。"])
h("3.6 统计与导出",2)
bullets(["统计以服务日期 service_date 分类汇总，不以订单创建时间作为业务统计日期。","订单按待确认（预约中）、已确认、已取消三种状态分别统计，并同时返回合计。","日期区间采用左闭右闭的自然日输入，数据库查询转换为明确边界；服务端统一 Asia/Shanghai。","管理员统计返回服务商 ID、名称、类别和各状态订单数；删除服务商仍可通过订单快照展示历史。","Excel 导出与列表使用同一筛选器；大数据量采用流式输出并限制最大时间范围/行数。"])

h("4. API 方案（建议基线）")
p("统一前缀 /api/v1；HTTP 状态码表达协议结果，业务错误码表达可处理原因；统一响应包含 code、message、data、requestId。")
table(["域","方法与路径","用途"],[
    ["微信身份","POST /app/wechat/session","一次性 code 静默交换 openid；不建立登录态"],
    ["大学","GET/POST /admin/universities；GET/PUT/DELETE /admin/universities/{id}","高校 CRUD"],
    ["服务商","GET/POST /admin/businesses；GET/PUT/DELETE /admin/businesses/{id}","服务商 CRUD、筛选"],
    ["附属数据","/admin/businesses/{id}/cars|dishes|rooms","按服务商类型维护子资源"],
    ["地图","GET /app/map-overview?type=","高校区域、点位、分类计数"],
    ["详情","GET /app/businesses/{id}","按类型返回服务商完整详情"],
    ["预约","POST /app/orders；GET /app/orders；GET /app/orders/{id}","创建与查询本人订单"],
    ["取消","POST /app/orders/{id}/cancel","条件状态变更"],
    ["管理订单","GET /admin/orders；GET /admin/orders/{id}","筛选与详情"],
    ["处理订单","POST /admin/orders/{id}/confirm|cancel","确认或取消待确认预约；管理员取消需原因"],
    ["统计/导出","GET /admin/statistics/overview；GET /admin/orders/export","统计与 Excel"],
    ["文件","POST /admin/files/images","受控图片上传，返回资源 ID/URL"],
],[1250,4850,3260])
h("4.1 通用约定",2)
bullets(["分页：page 从 1 开始，pageSize 默认 20、最大 100；稳定排序 createTime DESC, id DESC。","日期：业务日期 YYYY-MM-DD；时间戳 ISO-8601，存储与传输策略统一（建议数据库 UTC、展示上海时区）。","金额：decimal，不用浮点；枚举在服务端集中定义并在 OpenAPI 中声明。","敏感信息：日志与导出对手机号按权限控制并默认脱敏；错误响应不返回堆栈/SQL。","幂等：创建订单接收 Idempotency-Key，用户维度唯一并缓存/落库结果。"])

h("5. 数据库设计（推荐修订版）")
callout("设计原则","核心筛选、关联、约束和统计字段结构化；JSON 只容纳真正易变且无需检索的扩展信息。所有核心表包含 created_at、updated_at、deleted_at（或 deleted 标志），必要表增加 version。")
table(["表","用途","关键字段/关系","关键索引或约束"],[
    ["V001 遗留","已取消的 sys_admin、merchant_account","目标结构删除两表；使用后续 Flyway 迁移，不直接修改已执行 V001","迁移前检查数据；空库与升级路径均须验证"],
    ["wx_user","openid 映射（非登录系统）","id, openid, created_at, last_seen_at；可由首次订单/查询按需记录","UK(openid)"],
    ["university","高校区域","name, intro, polygon_json, status","IDX(status, deleted_at)"],
    ["business","服务商主表","name, address, lng, lat, type, intro, status, food_*","IDX(type,status,deleted_at)；经纬度 CHECK"],
    ["business_travel_car","车辆","business_id, model, seat_num, image_resource_id, status","FK business；IDX(business_id,status)"],
    ["business_food_dish","菜品","business_id, name, image_resource_id, sort_no","FK business；IDX(business_id)"],
    ["business_hotel_room","房型","business_id, name, bed_spec, description, image_resource_id, status","FK business；IDX(business_id,status)"],
    ["file_resource","文件资源","storage_key, public_url, mime, size, sha256, status","UK(storage_key)"],
    ["reserve_order","预约主表","order_no, user_id, business_id, contact_name/phone, people_num, service_date, service_type, status, type-specific fields, snapshots, cancel_source/reason, version","UK(order_no)；IDX(user_id,created_at)；IDX(service_date,status,type)；IDX(business_id,service_date)"],
    ["order_status_log","状态审计","order_id, from_status, to_status, operator_type/id, reason, created_at","IDX(order_id,created_at)"],
    ["api_idempotency","幂等记录","user_id, idempotency_key, request_hash, response_ref, expires_at","UK(user_id,key)"],
],[1450,1550,3650,2710])
h("5.1 reserve_order 关键字段建议",2)
table(["字段组","建议字段","说明"],[
    ["归属","openid NOT NULL；可选 user_id","openid 是前端传入的用户辨别字段；wx_user 仅做内部映射，不承担登录认证"],
    ["公共表单","contact_name, contact_phone, people_num, service_date","三类统一必填；service_date 是统计日期"],
    ["服务快照","business_id, business_name_snapshot, service_type","服务商名称和服务类别随单快照"],
    ["选择项","car_id/car_spec_snapshot/car_quantity；room_id/room_spec_snapshot/room_quantity；meal_period；service_mode","按类型校验必填性；数量必须为正整数"],
    ["历史快照","business_name_snapshot, option_snapshot_json","保证历史订单可读，不随服务商资料变化"],
    ["并发","version","订单状态乐观锁；也可使用 WHERE status=1 的条件更新"],
    ["审计","confirmed_at, cancelled_at, cancel_source, cancel_reason","区分 USER/ADMIN 取消；状态日志保留完整轨迹"],
],[1500,3800,4060])
h("5.2 删除与一致性策略",2)
bullets(["高校/服务商：逻辑删除；有历史订单的服务商禁止物理删除。","子资源：无订单引用可删除；被订单引用后逻辑删除/禁用，历史由快照兜底。","订单：业务数据不提供删除接口；必要时仅做合规匿名化。","服务商类型与子表一致性在服务层事务内校验；数据库外键保障存在性。","订单创建、状态更新、状态日志写入在同一事务完成。"])

h("6. 技术方案")
h("6.1 推荐技术栈",2)
table(["层","建议"],[
    ["运行时","JDK 1.8（Java 8），Spring Boot 2.7.x；固定在兼容 Java 8 的最终 2.7 版本"],
    ["Web/校验","Spring MVC, javax.validation, springdoc-openapi 1.x（与 Spring Boot 2.7 兼容）"],
    ["持久化","MyBatis-Plus 或 Spring Data JPA 二选一；Flyway 管理迁移"],
    ["身份边界","后端不引入账号密码、Session、JWT 或令牌；仅提供微信 code2Session 到 openid 的静默交换"],
    ["数据库","MySQL Community Server 9.0.1（本机服务 MySQL90，端口 3306），utf8mb4，严格 SQL mode"],
    ["文件","首期使用服务器本地受控目录 + Nginx（现有约 40GB）；通过 StorageService 抽象保留切换云存储能力"],
    ["文档测试","OpenAPI 3、JUnit 5、Testcontainers/MySQL、接口集合"],
    ["可观测性","Actuator 健康检查、结构化日志、requestId、慢 SQL 监控"],
],[1700,7660])
h("6.2 逻辑架构",2)
bullets(["Controller：协议适配、请求上下文、参数校验，不承载业务事务或登录认证。","Application/Service：用例编排、事务、业务规则与状态机。","Domain：枚举、规则、领域异常、订单状态转换。","Repository/Mapper：数据库访问；避免 Controller 直接调用 Mapper。","Integration：文件存储、Excel、微信 code2Session 等外部能力，使用接口隔离。","Query：地图聚合、管理员统计与导出使用专用查询。"])
h("6.3 安全与隐私",2)
bullets(["后端不保存或校验任何登录账号、密码、Session、JWT 或令牌；不返回或持久化微信 session_key。","微信 AppSecret 仅通过服务端环境变量配置，不进入 uni-app、通用配置或日志。","上传白名单 JPEG/PNG/WebP，校验真实 MIME、大小、随机文件名，阻断路径穿越和脚本文件。","SQL 参数化；CORS 精确配置管理员后台域名。","手机号属于个人信息：最小采集、传输 HTTPS、日志脱敏、导出审计、明确保留期限。","订单接口仍接收客户端 openid，未绑定服务端会话，不得描述为完整强认证。","接口限流：微信身份交换、管理员写操作、下单和导出重点保护；/admin/** 通过 Nginx 限制访问来源。"])
h("6.4 缓存策略",2)
p("MVP 不强制引入 Redis。地图聚合可先使用本机短 TTL 缓存或 HTTP ETag；高校/服务商启停及编辑后主动失效。图片首期写入服务器独立数据目录，数据库只保存相对存储键和访问 URL；必须设置容量告警、定期备份和上传限制。若容量或多实例需求增长，再迁移云对象存储。")

h("7. 推荐工程目录")
p("按单体分层架构起步，模块边界清晰，后续可演进，不建议 MVP 直接拆微服务。")
table(["目录","职责"],[
    ["src/main/java/.../common","统一响应、错误码、异常、分页、工具、基础配置"],
    [".../module/admin","平台管理员、内容、订单与全局管理用例"],
    [".../module/university","高校 controller/service/domain/repository/dto"],
    [".../module/business","服务商及 car/dish/room 子资源"],
    [".../module/order","预约创建、查询、状态机、状态日志"],
    [".../module/statistics","统计查询与导出"],
    [".../integration/storage","本地/对象存储抽象与实现"],
    ["src/main/resources/db/migration","Flyway V001__baseline.sql 等迁移"],
    ["src/main/resources/application*.yml","公共配置与环境覆盖；不提交真实密钥"],
    ["src/test/java","单元、集成、权限和状态机测试"],
    ["deploy/","Nginx 模板、Windows 服务脚本、部署说明（密钥除外）"],
],[3100,6260])

h("8. 非功能需求与验收")
table(["类别","要求"],[
    ["可靠性","事务保证订单与状态日志一致；数据库每日备份并定期恢复演练。"],
    ["性能","分页查询必须有索引；地图接口轻量化；避免附属数据 N+1；导出流式处理。"],
    ["兼容性","API 版本化；枚举新增向后兼容；数据库迁移可重复验证、不可手工漂移。"],
    ["日志","关键写操作记录请求来源、目标、结果与 requestId；禁止记录完整手机号/openid。"],
    ["错误处理","参数错误、资源不存在、状态冲突、外部服务失败分别编码。"],
    ["测试","核心服务单测；MySQL 集成测试；权限越权、重复下单、并发状态更新、边界日期必测。"],
    ["可用性","/actuator/health；服务异常可定位；上线配置与回滚步骤文档化。"],
],[1600,7760])
h("8.1 最低验收用例",2)
nums(["微信临时 code 可静默交换 openid，响应不含 session_key；无账号密码、Session、JWT、管理端令牌、商家账号或 /merchant/**。","生产环境缺少微信 AppID/AppSecret 时启动失败，日志不记录完整 code、openid 或密钥。","/admin/** 不做后端登录认证；Nginx 网络边界、HTTPS、限流与审计配置经过验证。","同一待确认订单被管理员确认与用户取消并发操作时，只有一个成功。","确认后用户取消必须失败；待确认订单允许对应 openid 取消。","禁用服务商不出现在地图/详情，且不能新建预约；历史订单仍可查看。","出行车辆/住宿房间必须属于对应服务商，车辆数量和房间数量为正整数。","管理员取消未填写原因时失败；非待确认状态重复处理失败。","按 service_date 和三种状态统计，管理员统计与订单明细一致。","上传伪装脚本、超大文件、非法 MIME 或路径穿越文件名均被拒绝。"])

h("9. 风险清单与应对")
table(["等级","风险","影响","应对"],[
    ["高","订单接口仍信任客户端 openid","可重放或篡改，静默交换不等于完整授权","全程 HTTPS、交换与订单接口限流、日志脱敏；后续可升级服务端会话绑定"],
    ["高","已取消商家角色仍残留在 V001 与代码","产生无效账号、路由、枚举和权限分支","新增 Flyway 迁移并清理 merchant_account、MERCHANT、/merchant/** 和相关代码测试"],
    ["高","预约字段放在 extra_json","筛选统计困难、校验弱、索引不可用","核心日期/人数/选择项结构化；JSON 仅扩展"],
    ["高","订单与资料无快照","服务商修改/删除后历史失真","订单保存名称与选择项快照；逻辑删除"],
    ["高","并发状态更新未定义","确认与取消互相覆盖","状态机 + 条件更新/version + 状态日志"],
    ["高","现有前端 mock 字段与后端口径不一致","联调时出现价格、金额、状态枚举或子资源字段缺失","先冻结映射；未确认字段不入库；由前端负责人按 OpenAPI 适配"],
    ["高","源文档对个人主体与备案结论过于绝对","小程序审核/上线受阻","上线前按实际服务类目、主体、域名与微信规则复核"],
    ["中","Windows 单机 40GB 文件存储","磁盘耗尽、备份与迁移困难","独立数据目录；格式/大小限制；容量告警与备份；StorageService 支持后续云存储"],
    ["中","手工录入多边形 JSON","坐标顺序/格式错误导致地图异常","结构校验、预览、点数限制、错误提示；后续地图框选"],
    ["低","统计实现偏离已确认口径","前后台数据争议","统一按 service_date；分别统计待确认、已确认、已取消"],
    ["中","直接放行 8080","绕过 Nginx/HTTPS 暴露后端","公网仅开放 80/443；8080 绑定内网/本机并由 Nginx 代理"],
    ["低","餐饮邻近高校使用自由文本","难筛选、易脏数据","MVP 可文本；若需要关联查询改 business_university_rel"],
],[700,2450,2750,3460])

h("10. 开发顺序与里程碑")
table(["顺序","交付物","前置/退出条件"],[
    ["01A-01B. 范围纠正","两角色契约；商家角色数据库与代码遗留清单","PRD/AGENTS/OpenAPI 一致；迁移方案明确"],
    ["02A. 登录遗留设计","盘点管理员/商家登录、密码、Session/JWT、令牌与安全规则","PRD/AGENTS/OpenAPI 无后端登录开发项"],
    ["02B. 遗留清理","Flyway 清理 sys_admin、merchant_account、MERCHANT；删除全部登录认证与 /merchant/** 分支","MySQL 升级路径、源码扫描和无登录契约测试通过"],
    ["03A-03D. 内容后端","高校、服务商、车辆/房型/菜品、文件 API 与联调契约","CRUD、归属、文件安全和逻辑删除测试通过；交付前端映射"],
    ["04A-04B. 管理员订单后端","订单筛选详情、确认取消和状态机 API","并发、终态、ADMIN 日志测试通过；交付现有订单页所需契约"],
    ["05A-05B. 统计导出后端","service_date 统计与 Excel 接口","统计与明细、导出筛选一致；交付下载说明"],
    ["06A-06B. 公开查询后端","地图聚合和服务商详情 API","仅返回启用数据；三类详情契约可供小程序渲染"],
    ["07A. 预约创建后端","幂等建单 API","快照、资源归属、并发防重通过；交付三类请求示例"],
    ["08A. 用户订单后端","openid 查询、详情和取消 API","仅待确认可取消；归属测试通过；交付状态映射"],
    ["09A-09C. 验收部署","自动化回归、MySQL/Flyway、Nginx/HTTPS、备份回滚","全部验收用例通过；不暴露 8080"],
],[1250,3000,5110])

h("11. 已确认决策与剩余问题")
table(["状态","事项","结论"],[
    ["已确认","小程序身份","uni-app 启动时静默取 code；Java 后端仅交换并返回 openid，不返回 session_key、不建立 Session/JWT"],
    ["已确认","系统角色","只有小程序用户和平台管理员；服务商只是管理员维护的业务数据"],
    ["已确认","后台范围","只建设平台管理员后台，不建设商家账号、商家后台、商家登录或 /merchant/**"],
    ["已确认","开发职责","本计划只实现后端、数据库、契约、测试与部署；管理员后台和小程序前端由其他开发人员完成"],
    ["已确认","现有管理员前端","已存在登录、仪表盘、服务商、高校和订单页面，目前使用 mock；后端按契约交付，不直接修改前端"],
    ["已确认","订单状态","仅待确认（预约中）、已确认、已取消；待确认可由用户取消，管理员确认后用户不可取消"],
    ["已确认","统计口径","按服务日期分类，并按三种订单状态分别汇总"],
    ["已确认","公用表单","联系人姓名、电话、人数、服务日期、服务商名称、服务类别"],
    ["已确认","分类表单","出行：车辆规格×数量、包天/往返；住宿：房间规格×数量；餐饮：早/中/晚餐"],
    ["已确认","图片存储","首期使用现有服务器约 40GB 空间；保留云存储适配层"],
    ["已确认","Java 基线","JDK 1.8 + Spring Boot 2.7.x"],
    ["P1","管理员接口允许哪些网络来源访问？","后端不认证；上线前必须确定 Nginx IP/VPN/内网访问边界"],
    ["P1","高校与服务商是否存在关联，一个服务商可邻近多个高校吗？","当前无强关联；若需按高校筛选，增加多对多关系"],
    ["P1","删除是逻辑删除还是物理删除？历史订单保存多久？","核心业务表逻辑删除；订单保留期需确认"],
    ["P1","预约是否需要防重复规则（同用户/服务商/日期）？","技术幂等必须做；业务重复是否禁止待确认"],
    ["P2","地图首屏数据规模预计多少（高校/服务商/多边形点数）？","规模不明，先轻量聚合 + 缓存/ETag"],
    ["P2","Excel 导出上限、字段与手机号权限？","默认最大 10,000 行、后台可见完整手机号并记录导出审计，需确认"],
    ["P2","是否接受对源文档部署建议调整为公网仅 80/443？","强烈建议接受，8080 不直接暴露公网"],
],[1000,4300,4060])
callout("推进建议","本次确认内容已可支持数据库与 OpenAPI 主体设计。剩余 P1/P2 可先采用本文默认值推进；若后续改变高校关联或删除策略，需要同步调整迁移脚本。",GOLD)

doc.add_page_break()
h("12. 需求追踪与版本说明")
table(["来源内容","本文处理"],[
    ["高校区域、三类服务商、预约、平台管理员后台、统计","纳入后端 MVP；前端页面由其他开发人员负责，本计划只保留接口契约与联调交付"],
    ["现有 TripAgency-web / TripAgency-uniapp","作为接口需求和字段映射依据；不在后端窗口中修改前端代码"],
    ["原稿与 V001 数据表","保留业务意图；保留 wx_user、file_resource、order_status_log、api_idempotency；merchant_account 作为取消角色遗留通过后续迁移清理"],
    ["无支付、只有管理员后台、手工坐标","作为明确版本约束；服务商没有账号或独立后台"],
    ["服务通知、评价、支付、地图框选、司机端","列为后续迭代，不进入当前实现"],
    ["个人主体/备案/Windows 适配结论","视为源文档陈述，不作为未经核验的合规保证"],
],[3100,6260])
# keep table rows together where practical; repeat headers
for t in doc.tables:
    trPr=t.rows[0]._tr.get_or_add_trPr(); rep=OxmlElement("w:tblHeader"); rep.set(qn("w:val"),"true"); trPr.append(rep)
    for row in t.rows:
        trPr=row._tr.get_or_add_trPr(); cant=OxmlElement("w:cantSplit"); trPr.append(cant)

doc.core_properties.title="地图商家预约系统后端与数据库PRD"
doc.core_properties.subject="后端开发、数据库设计、接口与实施计划"
doc.core_properties.author="Codex"
doc.save(OUT)
print(OUT)
